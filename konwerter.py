import os
from flask import Flask, render_template, request, redirect, url_for, send_from_directory
import fitz  # PyMuPDF
from docx import Document

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'

# Ensure the upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def pdf_to_docx(pdf_file, docx_file):
    try:
        doc = Document()
        pdf_document = fitz.open(pdf_file)
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text()
            doc.add_paragraph(text)
        doc.save(docx_file)
        pdf_document.close()
        return True, None
    except Exception as e:
        return False, str(e)

def docx_to_pdf(docx_file, pdf_file):
    try:
        doc = Document(docx_file)
        pdf_document = fitz.open()
        for paragraph in doc.paragraphs:
            page = pdf_document.new_page()
            page.insert_text((72, 72), paragraph.text, fontsize=12)
        pdf_document.save(pdf_file)
        pdf_document.close()
        return True, None
    except Exception as e:
        return False, str(e)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/konwertuj', methods=['POST'])
def konwertuj():
    format = request.form['format']
    plik = request.files['plik']
    
    if plik:
        filename = os.path.join(app.config['UPLOAD_FOLDER'], plik.filename)
        plik.save(filename)
        base_filename, file_extension = os.path.splitext(plik.filename)
        
        if format == 'pdf' and file_extension.lower() == '.docx':
            result_file = os.path.join(app.config['UPLOAD_FOLDER'], f"{base_filename}.pdf")
            success, error = docx_to_pdf(filename, result_file)
        elif format == 'docx' and file_extension.lower() == '.pdf':
            result_file = os.path.join(app.config['UPLOAD_FOLDER'], f"{base_filename}.docx")
            success, error = pdf_to_docx(filename, result_file)
        else:
            return "Nieprawidłowy format pliku do konwersji", 400
        
        if success:
            return send_from_directory(app.config['UPLOAD_FOLDER'], os.path.basename(result_file), as_attachment=True)
        else:
            return f"Wystąpił błąd: {error}", 500
    return "Nie wybrano pliku", 400

if __name__ == '__main__':
    app.run(debug=True)
